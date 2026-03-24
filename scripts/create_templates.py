"""One-time script to generate contract DOCX templates and test fixture.

Run: python scripts/create_templates.py
Requires: pip install python-docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return p


def build_production_template(doc, group_label: str):
    """Build a realistic contract template for one apartment group."""
    add_heading(doc, f"ДОГОВОР АРЕНДЫ ЖИЛОГО ПОМЕЩЕНИЯ № {{{{ contract_number }}}}")
    doc.add_paragraph()

    doc.add_paragraph("г. Москва                                    {{ contract_date }}")
    doc.add_paragraph()

    doc.add_paragraph(
        "Арендодатель, с одной стороны, и Арендатор — {{ tenant_full_name }}, "
        "дата рождения: {{ tenant_dob }}, место рождения: {{ tenant_birthplace }}, "
        "пол: {{ tenant_gender }}, адрес регистрации: {{ tenant_address }}, "
        "паспорт: серия {{ passport_series }} номер {{ passport_number }}, "
        "выдан {{ passport_issued_date }} — {{ passport_issued_by }}, "
        "код подразделения {{ passport_division_code }}, "
        "тел.: {{ tenant_phone }}, email: {{ tenant_email }}, "
        "с другой стороны, заключили настоящий договор."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        f"1. Объект аренды: квартира группы {group_label}."
    )
    doc.add_paragraph(
        "2. Дата начала аренды (Акт приёма-передачи): {{ act_date }}."
    )
    doc.add_paragraph(
        "3. Ежемесячная арендная плата составляет {{ monthly_amount }} руб."
    )
    doc.add_paragraph(
        "4. Депозит составляет {{ deposit_amount }} руб."
    )
    doc.add_paragraph()

    # Conditional deposit section — each {%p %} tag must be in its own paragraph
    p_if = doc.add_paragraph()
    p_if.add_run("{%p if deposit_split %}")

    doc.add_paragraph(
        "Депозит вносится в два платежа: {{ deposit_half }} руб. + {{ deposit_half }} руб."
    )

    p_endif1 = doc.add_paragraph()
    p_endif1.add_run("{%p endif %}")

    p_if2 = doc.add_paragraph()
    p_if2.add_run("{%p if not deposit_split %}")

    doc.add_paragraph(
        "Депозит вносится единовременно в размере {{ deposit_amount }} руб."
    )

    p_endif2 = doc.add_paragraph()
    p_endif2.add_run("{%p endif %}")

    doc.add_paragraph()
    doc.add_paragraph("Арендодатель: ___________________")
    doc.add_paragraph("Арендатор: ___________________   {{ tenant_full_name }}")


def build_test_fixture(doc):
    """Minimal template covering all context dict keys — used only in tests."""
    doc.add_paragraph("Договор № {{ contract_number }}")
    doc.add_paragraph("Дата: {{ contract_date }}")
    doc.add_paragraph("Акт: {{ act_date }}")
    doc.add_paragraph("Арендатор: {{ tenant_full_name }}")
    doc.add_paragraph("ДР: {{ tenant_dob }}")
    doc.add_paragraph("Место рождения: {{ tenant_birthplace }}")
    doc.add_paragraph("Пол: {{ tenant_gender }}")
    doc.add_paragraph("Адрес: {{ tenant_address }}")
    doc.add_paragraph("Паспорт: {{ passport_series }} {{ passport_number }}")
    doc.add_paragraph("Выдан: {{ passport_issued_date }} — {{ passport_issued_by }}")
    doc.add_paragraph("Код: {{ passport_division_code }}")
    doc.add_paragraph("Тел.: {{ tenant_phone }}")
    doc.add_paragraph("Email: {{ tenant_email }}")
    doc.add_paragraph("Аренда: {{ monthly_amount }} руб.")
    doc.add_paragraph("Депозит: {{ deposit_amount }} руб.")
    doc.add_paragraph("Половина: {{ deposit_half }} руб.")

    p_if = doc.add_paragraph()
    p_if.add_run("{%p if deposit_split %}")
    doc.add_paragraph("SPLIT: {{ deposit_half }} + {{ deposit_half }}")
    p_endif1 = doc.add_paragraph()
    p_endif1.add_run("{%p endif %}")

    p_if2 = doc.add_paragraph()
    p_if2.add_run("{%p if not deposit_split %}")
    doc.add_paragraph("LUMP: {{ deposit_amount }}")
    p_endif2 = doc.add_paragraph()
    p_endif2.add_run("{%p endif %}")


if __name__ == "__main__":
    root = Path(__file__).parent.parent

    # Production templates
    for group in ("Г39", "Г38"):
        out_path = root / "storage" / "templates" / group / "contract_template.docx"
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        build_production_template(doc, group)
        doc.save(str(out_path))
        print(f"Written: {out_path}")

    # Test fixture
    fixture_path = root / "tests" / "fixtures" / "contract_template_test.docx"
    fixture_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    build_test_fixture(doc)
    doc.save(str(fixture_path))
    print(f"Written: {fixture_path}")

    print("Done.")
