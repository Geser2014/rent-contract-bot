"""Document generation service — TXT template fill + LibreOffice PDF conversion.

Public API:
    generate_contract_number(group, apartment, contract_date) -> str
    generate_contract(data: ContractData, extra: dict) -> str  [async, returns PDF path]
"""
import asyncio
import json
import logging
import shutil
import subprocess
import sys
from pathlib import Path

import config
from models import ContractData

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Apartments config (fixed per-apartment data)
# ---------------------------------------------------------------------------
_APARTMENTS_FILE = Path(__file__).parent / "apartments.json"


def load_apartments() -> dict:
    """Load apartments.json with fixed per-apartment fields."""
    if _APARTMENTS_FILE.exists():
        with open(_APARTMENTS_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}


APARTMENTS_DATA = load_apartments()


def get_apartment_names(group: str) -> list[str]:
    """Return list of apartment names for a group (excludes _short metadata)."""
    return [k for k in APARTMENTS_DATA.get(group, {}).keys() if not k.startswith("_")]


def get_apartment_fixed_data(group: str, apartment: str) -> dict:
    """Return fixed data for an apartment (address, rooms, area, inventory)."""
    return APARTMENTS_DATA.get(group, {}).get(apartment, {})


# ---------------------------------------------------------------------------
# Number generation
# ---------------------------------------------------------------------------

def generate_contract_number(group: str, apartment: str, contract_date) -> str:
    """Generate contract number: П38/4/220209.

    Uses short group code (_short from apartments.json) and apartment
    contract_num for compact numbering. Date format: DDMMYY.
    """
    group_data = APARTMENTS_DATA.get(group, {})
    short_group = group_data.get("_short", group)
    apt_data = group_data.get(apartment, {})
    apt_num = apt_data.get("contract_num", apartment)
    date_str = contract_date.strftime("%d%m%y")
    return f"{short_group}/{apt_num}/{date_str}"


# ---------------------------------------------------------------------------
# Amount to words (Russian)
# ---------------------------------------------------------------------------

_UNITS = [
    "", "один", "два", "три", "четыре", "пять",
    "шесть", "семь", "восемь", "девять",
]
_UNITS_F = [
    "", "одна", "две", "три", "четыре", "пять",
    "шесть", "семь", "восемь", "девять",
]
_TEENS = [
    "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
    "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать",
]
_TENS = [
    "", "", "двадцать", "тридцать", "сорок", "пятьдесят",
    "шестьдесят", "семьдесят", "восемьдесят", "девяносто",
]
_HUNDREDS = [
    "", "сто", "двести", "триста", "четыреста", "пятьсот",
    "шестьсот", "семьсот", "восемьсот", "девятьсот",
]


def _int_to_words(n: int) -> str:
    """Convert integer to Russian words (up to 999 999 999)."""
    if n == 0:
        return "ноль"

    parts = []

    # millions
    millions = n // 1_000_000
    if millions:
        parts.append(_triplet_to_words(millions, feminine=False))
        if millions % 10 == 1 and millions % 100 != 11:
            parts.append("миллион")
        elif 2 <= millions % 10 <= 4 and not (12 <= millions % 100 <= 14):
            parts.append("миллиона")
        else:
            parts.append("миллионов")

    # thousands
    thousands = (n % 1_000_000) // 1000
    if thousands:
        parts.append(_triplet_to_words(thousands, feminine=True))
        if thousands % 10 == 1 and thousands % 100 != 11:
            parts.append("тысяча")
        elif 2 <= thousands % 10 <= 4 and not (12 <= thousands % 100 <= 14):
            parts.append("тысячи")
        else:
            parts.append("тысяч")

    # units
    remainder = n % 1000
    if remainder:
        parts.append(_triplet_to_words(remainder, feminine=False))

    return " ".join(parts)


def _triplet_to_words(n: int, feminine: bool) -> str:
    """Convert 1-999 to Russian words."""
    parts = []
    h = n // 100
    if h:
        parts.append(_HUNDREDS[h])
    t = (n % 100) // 10
    u = n % 10
    if t == 1:
        parts.append(_TEENS[u])
    else:
        if t:
            parts.append(_TENS[t])
        if u:
            units = _UNITS_F if feminine else _UNITS
            parts.append(units[u])
    return " ".join(parts)


def amount_to_words(amount: int) -> str:
    """Convert amount to Russian words: '50000' -> 'пятьдесят тысяч'."""
    return _int_to_words(amount)


# ---------------------------------------------------------------------------
# Day number to Russian ordinal word
# ---------------------------------------------------------------------------

_DAY_WORDS = {
    1: "первого", 2: "второго", 3: "третьего", 4: "четвертого", 5: "пятого",
    6: "шестого", 7: "седьмого", 8: "восьмого", 9: "девятого", 10: "десятого",
    11: "одиннадцатого", 12: "двенадцатого", 13: "тринадцатого",
    14: "четырнадцатого", 15: "пятнадцатого", 16: "шестнадцатого",
    17: "семнадцатого", 18: "восемнадцатого", 19: "девятнадцатого",
    20: "двадцатого", 21: "двадцать первого", 22: "двадцать второго",
    23: "двадцать третьего", 24: "двадцать четвертого", 25: "двадцать пятого",
    26: "двадцать шестого", 27: "двадцать седьмого", 28: "двадцать восьмого",
    29: "двадцать девятого", 30: "тридцатого", 31: "тридцать первого",
}


# ---------------------------------------------------------------------------
# Template fill (TXT-based, [PLACEHOLDER] format)
# ---------------------------------------------------------------------------

def _build_replacements(data: ContractData, extra: dict) -> dict[str, str]:
    """Build placeholder->value mapping for TXT template."""
    act_day = data.act_date.day
    monthly_int = int(data.monthly_amount)
    deposit_int = int(data.deposit_amount)

    # Deposit condition text
    if data.deposit_split:
        half = deposit_int // 2
        deposit_condition = (
            f"Обеспечительный платеж вносится в два этапа: "
            f"{half} ({amount_to_words(half)}) руб. при подписании договора, "
            f"{half} ({amount_to_words(half)}) руб. в течение 30 дней."
        )
    else:
        deposit_condition = ""

    # Fixed apartment data
    apt_data = get_apartment_fixed_data(data.group, data.apartment)

    replacements = {
        # Contract
        "[НОМЕР_ДОГОВОРА]": data.contract_number,
        "[ДАТА_ДОГОВОРА]": data.contract_date.strftime("%d.%m.%Y"),
        "[ДАТА_АКТА]": data.act_date.strftime("%d.%m.%Y"),

        # Tenant (from OCR)
        "[ФИО_АРЕНДАТОРА]": data.tenant_full_name,
        "[ПОЛ]": data.tenant_gender,
        "[ДАТА_РОЖДЕНИЯ]": data.tenant_dob.strftime("%d %B %Y").replace(
            "January", "января").replace("February", "февраля").replace(
            "March", "марта").replace("April", "апреля").replace(
            "May", "мая").replace("June", "июня").replace(
            "July", "июля").replace("August", "августа").replace(
            "September", "сентября").replace("October", "октября").replace(
            "November", "ноября").replace("December", "декабря"),
        "[МЕСТО_РОЖДЕНИЯ]": data.tenant_birthplace,
        "[СЕРИЯ_ПАСПОРТА]": data.passport_series,
        "[НОМЕР_ПАСПОРТА]": data.passport_number,
        "[КЕМ_ВЫДАН]": data.passport_issued_by,
        "[ДАТА_ВЫДАЧИ]": data.passport_issued_date.strftime("%d %B %Y").replace(
            "January", "января").replace("February", "февраля").replace(
            "March", "марта").replace("April", "апреля").replace(
            "May", "мая").replace("June", "июня").replace(
            "July", "июля").replace("August", "августа").replace(
            "September", "сентября").replace("October", "октября").replace(
            "November", "ноября").replace("December", "декабря"),
        "[КОД_ПОДРАЗДЕЛЕНИЯ]": data.passport_division_code,
        "[АДРЕС_РЕГИСТРАЦИИ]": data.tenant_address,
        "[ТЕЛЕФОН]": data.tenant_phone,
        "[EMAIL]": data.tenant_email,
        "[ТЕЛЕГРАМ]": extra.get("telegram", "___"),

        # Financial
        "[СУММА_АРЕНДЫ]": str(monthly_int),
        "[СУММА_АРЕНДЫ_ПРОПИСЬЮ]": amount_to_words(monthly_int),
        "[ДЕПОЗИТ]": str(deposit_int),
        "[ДЕПОЗИТ_ПРОПИСЬЮ]": amount_to_words(deposit_int),
        "[УСЛОВИЕ_ДЕПОЗИТ_НА_2_ЧАСТИ]": deposit_condition,
        "[ДЕНЬ_ОПЛАТЫ_ЦИФРА]": str(act_day),
        "[ДЕНЬ_ОПЛАТЫ_ПРОПИСЬЮ]": _DAY_WORDS.get(act_day, str(act_day)),

        # Fixed per-apartment
        "[АДРЕС_КВАРТИРЫ]": apt_data.get("address", "___"),
        "[КОЛИЧЕСТВО_КОМНАТ]": apt_data.get("rooms", "___"),
        "[ПЛОЩАДЬ]": apt_data.get("area", "___"),
        "[ОПИСЬ_ИМУЩЕСТВА]": apt_data.get("inventory", "___"),

        # Extra from dialog
        "[СПИСОК_ПРОЖИВАЮЩИХ]": extra.get("residents", "___"),
        "[СРОК_ДОГОВОРА]": extra.get("contract_duration", "___"),
        "[ДОП_УСЛОВИЯ]": extra.get("extra_conditions", "Нет"),
    }
    return replacements


def _fill_txt_template(template_path: Path, replacements: dict[str, str]) -> str:
    """Read TXT template and replace all [PLACEHOLDER] tokens."""
    text = template_path.read_text(encoding="utf-8")
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, value)
    return text


def _fill_docx_template(template_path: Path, replacements: dict[str, str], output_path: Path) -> Path:
    """Fill DOCX template with [PLACEHOLDER] replacement. Returns path to filled DOCX."""
    from docx import Document

    doc = Document(str(template_path))

    def replace_in_paragraph(paragraph):
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# LibreOffice PDF conversion
# ---------------------------------------------------------------------------

def _find_libreoffice() -> str:
    """Find LibreOffice executable for the current platform."""
    if sys.platform == "win32":
        for candidate in [
            Path("C:/Program Files/LibreOffice/program/soffice.exe"),
            Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        ]:
            if candidate.exists():
                return str(candidate)
    return shutil.which("libreoffice") or "libreoffice"


def _convert_to_pdf(src_path: Path, out_dir: Path) -> Path:
    """Convert TXT or DOCX to PDF via LibreOffice headless."""
    out_dir.mkdir(parents=True, exist_ok=True)

    lo_bin = _find_libreoffice()
    result = subprocess.run(
        [
            lo_bin,
            "--headless",
            "--norestore",
            "--nofirststartwizard",
            "--convert-to", "pdf",
            "--outdir", str(out_dir.resolve()),
            str(src_path.resolve()),
        ],
        capture_output=True,
        timeout=60,
    )

    expected_pdf = out_dir / (src_path.stem + ".pdf")
    if not expected_pdf.exists():
        raise RuntimeError(
            f"LibreOffice conversion produced no output. "
            f"returncode={result.returncode} "
            f"stderr={result.stderr.decode('utf-8', errors='replace')!r}"
        )
    return expected_pdf


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

async def generate_contract(data: ContractData, extra: dict | None = None) -> str:
    """Fill template (TXT or DOCX) and convert to PDF. Returns absolute PDF path.

    Args:
        data: ContractData with all contract fields
        extra: dict with additional fields (telegram, residents, contract_duration, extra_conditions)
    """
    if extra is None:
        extra = {}

    # Find template — prefer .docx, fallback to .txt
    docx_path = config.TEMPLATES_DIR / data.group / f"{data.apartment}.docx"
    txt_path = config.TEMPLATES_DIR / data.group / f"{data.apartment}.txt"

    if docx_path.exists():
        template_path = docx_path
        template_type = "docx"
    elif txt_path.exists():
        template_path = txt_path
        template_type = "txt"
    else:
        raise FileNotFoundError(
            f"Template not found: {docx_path} or {txt_path}"
        )

    # Build replacements and fill
    replacements = _build_replacements(data, extra)
    safe_name = data.contract_number.replace("/", "_")
    out_dir = config.CONTRACTS_DIR / data.group / data.apartment
    out_dir.mkdir(parents=True, exist_ok=True)

    if template_type == "docx":
        filled_path = out_dir / f"{safe_name}.docx"
        _fill_docx_template(template_path, replacements, filled_path)
    else:
        filled_path = out_dir / f"{safe_name}.txt"
        filled_text = _fill_txt_template(template_path, replacements)
        filled_path.write_text(filled_text, encoding="utf-8")

    logger.info("Template filled (%s): %s", template_type, filled_path)

    # Convert to PDF via LibreOffice
    pdf_path = await asyncio.to_thread(_convert_to_pdf, filled_path, out_dir)
    logger.info("PDF generated: %s", pdf_path)
    return str(pdf_path)
